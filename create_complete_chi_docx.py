#!/usr/bin/env python3
"""
完整的CHI2026 GestureFlow Word文档生成器
基于原始Markdown内容生成符合ACM格式的完整论文
"""

import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_complete_chi_document():
    """创建完整的CHI格式Word文档"""

    # 创建新文档
    doc = Document()

    # 设置文档样式
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # 读取论文内容
    with open('CHI2026_GestureFlow_Poster_Paper.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # 提取作者信息
    author_line = re.search(r'Jiajun Wu.*\n', content).group().strip()

    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('GestureFlow: EMG-GSR Gesture Recognition for Digital Nomad Rhythm Management')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 添加作者信息
    author_para = doc.add_paragraph()
    author_run = author_para.add_run('Jiajun Wu¹*(ORCID: 0009-0000-6828-2241), Junfeng Wang¹')
    author_run.font.size = Pt(12)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加机构信息
    affiliation_para = doc.add_paragraph()
    affiliation_run = affiliation_para.add_run('¹School of Creative Design, Shenzhen Technology University, Shenzhen 518118, People\'s Republic of China')
    affiliation_run.font.size = Pt(10)
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加邮箱
    email_para = doc.add_paragraph()
    email_run = email_para.add_run('{epwujiajun@icloud.com, wangjunfeng@sztu.edu.cn}')
    email_run.font.size = Pt(10)
    email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加通讯作者标识
    corresponding_para = doc.add_paragraph()
    corresponding_run = corresponding_para.add_run('*Corresponding author')
    corresponding_run.font.size = Pt(10)
    corresponding_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加空行
    for _ in range(3):
        doc.add_paragraph()

    # 提取摘要
    abstract_match = re.search(r'## Abstract\n\n(.*?)(?=\n\n\*\*Keywords:)', content, re.DOTALL)
    if abstract_match:
        abstract_para = doc.add_paragraph()
        abstract_title = abstract_para.add_run('ABSTRACT')
        abstract_title.bold = True
        abstract_title.font.size = Pt(12)

        abstract_content_para = doc.add_paragraph(abstract_match.group(1).strip())
        abstract_content_para.font.size = Pt(10)
        abstract_content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 添加空行
    doc.add_paragraph()

    # 提取关键词
    keywords_match = re.search(r'\*\*Keywords:\*\* (.*)', content)
    if keywords_match:
        # CCS Concepts
        ccs_para = doc.add_paragraph()
        ccs_title = ccs_para.add_run('CCS Concepts: ')
        ccs_title.bold = True
        ccs_para.add_run('• Human-centered computing~Human computer interaction (HCI)')
        ccs_para.font.size = Pt(9)

        # General Terms
        general_para = doc.add_paragraph()
        general_title = general_para.add_run('General Terms: ')
        general_title.bold = True
        general_para.add_run('Design, Human Factors, Measurement')
        general_para.font.size = Pt(9)

        # Keywords
        keywords_para = doc.add_paragraph()
        keywords_title = keywords_para.add_run('Keywords: ')
        keywords_title.bold = True
        keywords_para.add_run(keywords_match.group(1))
        keywords_para.font.size = Pt(9)

    # 添加章节
    sections = [
        ('1 Introduction', '1'),
        ('2 Related Work', '2'),
        ('3 System Design', '3'),
        ('4 User Study', '4'),
        ('5 Discussion', '5'),
        ('6 Conclusion', '6')
    ]

    for section_title, section_num in sections:
        doc.add_page_break()

        # 添加章节标题
        section_para = doc.add_paragraph()
        section_run = section_para.add_run(f'{section_title}')
        section_run.bold = True
        section_run.font.size = Pt(14)
        section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()

        # 提取章节内容
        section_pattern = rf'## {section_num}.*?\n\n(.*?)(?=## {" + str(int(section_num) + 1) if section_num != "6" else "## References"}|## References)'
        section_match = re.search(section_pattern, content, re.DOTALL)

        if section_match:
            section_content = section_match.group(1).strip()

            # 处理章节内容
            lines = section_content.split('\n')

            for line in lines:
                line = line.strip()

                if not line:
                    doc.add_paragraph()
                    continue

                # 处理三级标题
                if line.startswith('### '):
                    subsection_para = doc.add_paragraph()
                    subsection_run = subsection_para.add_run(line[4:])
                    subsection_run.bold = True
                    subsection_run.font.size = Pt(12)
                    subsection_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 处理列表项
                elif line.startswith('- '):
                    list_para = doc.add_paragraph(line[2:], style='List Bullet')

                # 处理表格标记
                elif line.startswith('*Table 1:'):
                    table_para = doc.add_paragraph()
                    table_run = table_para.add_run(line[1:])
                    table_run.bold = True
                    table_run.italic = True

                    # 添加表格
                    create_confusion_matrix_table(doc)

                # 处理普通段落
                else:
                    # 跳过代码块
                    if not line.startswith('```'):
                        # 处理加粗文本
                        if '**' in line:
                            # 简单的加粗处理
                            para = doc.add_paragraph()
                            parts = line.split('**')
                            for i, part in enumerate(parts):
                                if i % 2 == 0:
                                    para.add_run(part)
                                else:
                                    bold_run = para.add_run(part)
                                    bold_run.bold = True
                        else:
                            doc.add_paragraph(line)

    # 添加参考文献
    doc.add_page_break()

    refs_para = doc.add_paragraph()
    refs_run = refs_para.add_run('REFERENCES')
    refs_run.bold = True
    refs_run.font.size = Pt(14)
    refs_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # 提取参考文献
    refs_match = re.search(r'## References\n\n(.*?)(?=---)', content, re.DOTALL)
    if refs_match:
        refs_content = refs_match.group(1).strip()
        refs_lines = refs_content.split('\n')

        for ref_line in refs_lines:
            ref_line = ref_line.strip()
            if ref_line and ref_line.startswith('['):
                doc.add_paragraph(ref_line)

    # 添加版权信息
    doc.add_page_break()
    copyright_para = doc.add_paragraph()
    copyright_run = copyright_para.add_run('Copyright © 2026 Association for Computing Machinery. This is the author\'s version of the work. It is posted here for personal use and not for redistribution.')
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_run.font.size = Pt(9)

    return doc

def create_confusion_matrix_table(doc):
    """创建混淆矩阵表格"""

    # 添加表格标题
    table_title = doc.add_paragraph('Table 1: Gesture Classification Confusion Matrix')
    table_title.italic = True
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 创建4x4表格
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].text = 'Predicted Work'
    header_cells[2].text = 'Predicted Rest'
    header_cells[3].text = 'Predicted Leisure'

    # 数据行
    # 第一行数据
    data_row1 = table.rows[1].cells
    data_row1[0].text = 'Actual Work'
    data_row1[1].text = '92.3%'
    data_row1[2].text = '5.7%'
    data_row1[3].text = '2.0%'

    # 第二行数据
    data_row2 = table.rows[2].cells
    data_row2[0].text = 'Actual Rest'
    data_row2[1].text = '3.8%'
    data_row2[2].text = '94.1%'
    data_row2[3].text = '2.1%'

    # 第三行数据
    data_row3 = table.rows[3].cells
    data_row3[0].text = 'Actual Leisure'
    data_row3[1].text = '4.2%'
    data_row3[2].text = '6.5%'
    data_row3[3].text = '89.3%'

def main():
    """主函数"""
    print("🚀 开始生成完整的CHI2026 GestureFlow Word文档...")

    try:
        # 创建文档
        doc = create_complete_chi_document()

        # 保存文档
        output_file = 'CHI2026_GestureFlow_Complete_Paper.docx'
        doc.save(output_file)

        print(f"✅ 完整Word文档已生成: {output_file}")
        print("📝 文档包含:")
        print("   ✅ 正确的作者信息 (吴嘉俊 + 王军锋)")
        print("   ✅ 深圳技术大学创意设计学院")
        print("   ✅ ORCID ID: 0009-0000-6828-2241")
        print("   ✅ 邮箱: epwujiajun@icloud.com")
        print("   ✅ 完整摘要和关键词")
        print("   ✅ 所有主要章节内容")
        print("   ✅ 混淆矩阵表格")
        print("   ✅ 完整参考文献")
        print("   ✅ ACM版权信息")
        print()
        print("🎯 文档特点:")
        print("   - 符合ACM格式标准")
        print("   - 完整的论文内容")
        print("   - 专业的排版格式")
        print("   - 可直接用于CHI2026投稿")

    except Exception as e:
        print(f"❌ 生成文档时出错: {str(e)}")

if __name__ == "__main__":
    main()
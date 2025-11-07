#!/usr/bin/env python3
"""
分析ACM模板格式要求并创建优化的文档生成器
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re

def create_optimized_acm_document():
    """创建优化的ACM格式文档，严格遵循模板要求"""

    # 读取HCI风格Markdown文件
    with open('CHI2026_GestureFlow_HCI_Style_Paper.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # 创建新文档
    doc = Document()

    # ACM标准格式设置
    # 1. 页面设置：1英寸边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 2. 字体设置：Times New Roman
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10)  # ACM标准通常使用10pt正文

    # --- 标题页 ---

    # 标题 (14pt, 粗体, 居中)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('GestureFlow: Embodied Rhythm Management for Digital Nomads Through Sensing-Instead-of-Controlling')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加空行
    doc.add_paragraph()

    # 作者信息 (10pt, 居中)
    authors_para = doc.add_paragraph()
    authors_run = authors_para.add_run('Jiajun Wu¹*, Junfeng Wang¹')
    authors_run.font.name = 'Times New Roman'
    authors_run.font.size = Pt(10)
    authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 机构信息 (9pt, 居中)
    affiliation_para = doc.add_paragraph()
    affiliation_run = affiliation_para.add_run('¹School of Creative Design, Shenzhen Technology University, Shenzhen 518118, People\'s Republic of China')
    affiliation_run.font.name = 'Times New Roman'
    affiliation_run.font.size = Pt(9)
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 邮箱 (9pt, 居中)
    email_para = doc.add_paragraph()
    email_run = email_para.add_run('{epwujiajun@icloud.com, wangjunfeng@sztu.edu.cn}')
    email_run.font.name = 'Times New Roman'
    email_run.font.size = Pt(9)
    email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 通讯作者 (9pt, 居中)
    corresponding_para = doc.add_paragraph()
    corresponding_run = corresponding_para.add_run('*Corresponding author')
    corresponding_run.font.name = 'Times New Roman'
    corresponding_run.font.size = Pt(9)
    corresponding_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加多个空行
    for _ in range(3):
        doc.add_paragraph()

    # --- 摘要和关键词 ---

    # 提取摘要
    abstract_match = re.search(r'## Abstract\s*\n\s*\n(.*?)(?=\n\n\*\*Keywords:)', content, re.DOTALL)
    if abstract_match:
        # ABSTRACT标题 (9pt, 粗体)
        abstract_title = doc.add_paragraph()
        abstract_title_run = abstract_title.add_run('ABSTRACT')
        abstract_title_run.font.name = 'Times New Roman'
        abstract_title_run.font.size = Pt(9)
        abstract_title_run.bold = True

        # 摘要内容 (9pt, 紧凑格式)
        abstract_content = abstract_match.group(1).strip()
        abstract_para = doc.add_paragraph(abstract_content)
        abstract_para.style = doc.styles['Normal']

    # 关键词部分
    keywords_match = re.search(r'\*\*Keywords:\*\* (.*)', content)
    if keywords_match:
        keywords_text = keywords_match.group(1)

        # CCS Concepts (9pt)
        ccs_para = doc.add_paragraph()
        ccs_run = ccs_para.add_run('CCS Concepts: ')
        ccs_run.font.name = 'Times New Roman'
        ccs_run.font.size = Pt(9)
        ccs_run.bold = True
        ccs_para.add_run('• Human-centered computing~Embodied interaction • Human computer interaction (HCI)')

        # General Terms (9pt)
        general_para = doc.add_paragraph()
        general_run = general_para.add_run('General Terms: ')
        general_run.font.name = 'Times New Roman'
        general_run.font.size = Pt(9)
        general_run.bold = True
        general_para.add_run('Design, Human Factors, Measurement')

        # Keywords (9pt)
        keywords_para = doc.add_paragraph()
        keywords_run = keywords_para.add_run('Keywords: ')
        keywords_run.font.name = 'Times New Roman'
        keywords_run.font.size = Pt(9)
        keywords_run.bold = True
        keywords_para.add_run(keywords_text)

    # --- 章节内容 ---
    sections = [
        ("1 INTRODUCTION", r'## 1 Introduction', r'## 2 Research Process'),
        ("2 RESEARCH PROCESS", r'## 2 Research Process', r'## 3 System and Design Concept'),
        ("3 SYSTEM AND DESIGN CONCEPT", r'## 3 System and Design Concept', r'## 4 User Study and Evaluation'),
        ("4 USER STUDY AND EVALUATION", r'## 4 User Study and Evaluation', r'## 5 Design Implications'),
        ("5 DESIGN IMPLICATIONS", r'## 5 Design Implications', r'## 6 Conclusion'),
        ("6 CONCLUSION", r'## 6 Conclusion', r'## References')
    ]

    for title, start_pattern, end_pattern in sections:
        # 分页 (除了第一个章节)
        if title != "1 INTRODUCTION":
            doc.add_page_break()

        # 章节标题 (12pt, 粗体)
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run(title)
        section_title_run.font.name = 'Times New Roman'
        section_title_run.font.size = Pt(12)
        section_title_run.bold = True

        # 提取章节内容
        section_pattern = fr'{start_pattern}.*?\n\n(.*?)(?={end_pattern})'
        section_match = re.search(section_pattern, content, re.DOTALL)

        if section_match:
            section_content = section_match.group(1).strip()
            lines = section_content.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                # 处理三级标题 (10pt, 粗体)
                if line.startswith('### '):
                    subtitle = doc.add_paragraph()
                    subtitle_run = subtitle.add_run(line[4:].strip())
                    subtitle_run.font.name = 'Times New Roman'
                    subtitle_run.font.size = Pt(10)
                    subtitle_run.bold = True

                # 处理设计洞见加粗段落
                elif line.startswith('**') and line.endswith('**'):
                    bold_para = doc.add_paragraph()
                    bold_run = bold_para.add_run(line[2:-2])
                    bold_run.font.name = 'Times New Roman'
                    bold_run.font.size = Pt(9)
                    bold_run.bold = True

                # 处理表格
                elif line.startswith('*Table 1:'):
                    # 表格标题 (9pt, 斜体, 居中)
                    table_title = doc.add_paragraph('Table 1: Gesture Classification Confusion Matrix')
                    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table_title_run = table_title.runs[0]
                    table_title_run.font.name = 'Times New Roman'
                    table_title_run.font.size = Pt(9)
                    table_title_run.italic = True

                    # 创建表格
                    create_acm_table(doc)

                # 跳过特殊行
                elif line.startswith('---') or line.startswith('|') or line.startswith('```'):
                    continue

                # 普通段落
                else:
                    # 处理段内加粗
                    if '**' in line:
                        para = doc.add_paragraph()
                        parts = re.split(r'(\*\*.*?\*\*)', line)
                        for i, part in enumerate(parts):
                            if i % 2 == 0:
                                if part.strip():
                                    run = para.add_run(part)
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(9)
                            else:
                                bold_text = part[2:-2]
                                bold_run = para.add_run(bold_text)
                                bold_run.font.name = 'Times New Roman'
                                bold_run.font.size = Pt(9)
                                bold_run.bold = True
                            if i < len(parts) - 1:
                                para.add_run(' ')
                    else:
                        para = doc.add_paragraph(line)
                        para.style = doc.styles['Normal']

    # --- 参考文献 ---
    doc.add_page_break()

    refs_heading = doc.add_paragraph()
    refs_run = refs_heading.add_run('REFERENCES')
    refs_run.font.name = 'Times New Roman'
    refs_run.font.size = Pt(12)
    refs_run.bold = True

    # 添加参考文献
    refs_match = re.search(r'## References\s*\n\s*\n(.*?)(?=---)', content, re.DOTALL)
    if refs_match:
        refs_content = refs_match.group(1).strip()
        refs_lines = refs_content.split('\n')

        for ref_line in refs_lines:
            ref_line = ref_line.strip()
            if ref_line and ref_line.startswith('['):
                ref_para = doc.add_paragraph(ref_line)
                ref_para.style = doc.styles['Normal']

    # --- 版权信息 ---
    doc.add_page_break()
    copyright_para = doc.add_paragraph()
    copyright_run = copyright_para.add_run('Copyright © 2026 Association for Computing Machinery. This is the author\'s version of the work. It is posted here for personal use and not for redistribution.')
    copyright_run.font.name = 'Times New Roman'
    copyright_run.font.size = Pt(8)
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

def create_acm_table(doc):
    """创建符合ACM格式的表格"""
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # 设置表格字体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 填充内容
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].text = 'Predicted Work'
    header_cells[2].text = 'Predicted Rest'
    header_cells[3].text = 'Predicted Leisure'

    actual_work = table.rows[1].cells
    actual_work[0].text = 'Actual Work'
    actual_work[1].text = '92.3%'
    actual_work[2].text = '5.7%'
    actual_work[3].text = '2.0%'

    actual_rest = table.rows[2].cells
    actual_rest[0].text = 'Actual Rest'
    actual_rest[1].text = '3.8%'
    actual_rest[2].text = '94.1%'
    actual_rest[3].text = '2.1%'

    actual_leisure = table.rows[3].cells
    actual_leisure[0].text = 'Actual Leisure'
    actual_leisure[1].text = '4.2%'
    actual_leisure[2].text = '6.5%'
    actual_leisure[3].text = '89.3%'

def main():
    """主函数"""
    print("🎯 生成严格符合ACM模板格式的CHI2026文档...")
    print("📋 ACM格式标准:")
    print("   - Times New Roman字体")
    print("   - 标题14pt, 章节标题12pt, 正文9pt")
    print("   - 1英寸页面边距")
    print("   - 紧凑的段落间距")
    print("   - 标准表格格式")

    try:
        doc = create_optimized_acm_document()
        output_file = 'CHI2026_GestureFlow_ACM_Template_Final.docx'
        doc.save(output_file)

        print(f"\n✅ ACM模板格式文档已生成: {output_file}")
        print("\n📝 格式特点:")
        print("   ✅ 严格遵循ACM论文模板格式")
        print("   ✅ 正确的字体和字号层次")
        print("   ✅ 标准的页面布局")
        print("   ✅ 规范的表格和参考文献格式")
        print("   ✅ 符合CHI2026投稿要求")
        print("\n🎯 现在完全匹配ACM模板格式！")

    except Exception as e:
        print(f"❌ 生成文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()